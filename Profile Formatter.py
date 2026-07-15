import streamlit as st
from docx import Document
from io import BytesIO
from openai import AzureOpenAI
import logging
from pathlib import Path

LOG_DIR = Path("logs")
LOG_DIR.mkdir(exist_ok=True)

LOG_FILE = LOG_DIR / "app.log"

logging.basicConfig(
    filename=str(LOG_FILE),
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s"
)

logger = logging.getLogger(__name__)

logger.info("Application started")

# -----------------------------
# Azure OpenAI Configuration
# -----------------------------
try:
    client = AzureOpenAI(
        api_key=st.secrets["AZURE_OPENAI_KEY"],
        azure_endpoint=st.secrets["AZURE_OPENAI_ENDPOINT"],
        api_version="2024-02-15-preview"
    )

    DEPLOYMENT_NAME = st.secrets["DEPLOYMENT_NAME"]

except Exception as e:
    st.error(f"Azure OpenAI configuration error: {e}")
    st.stop()


# -----------------------------
# Helper Functions
# -----------------------------
def read_docx(file):
    """Read uploaded Word document."""
    doc = Document(file)
    text = "\n".join(
        paragraph.text
        for paragraph in doc.paragraphs
        if paragraph.text.strip()
    )
    return text


def create_word_document(content):
    """Create downloadable Word document."""
    doc = Document()

    for line in content.split("\n"):
        doc.add_paragraph(line)

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    return output


def generate_profile(profile_text, template_text):
    """Generate formatted profile using Azure OpenAI."""

    prompt = f"""
You are an expert profile writer.

TASK:
Convert the profile information provided into the exact format,
structure, headings, and style used in the template.

TEMPLATE:
--------------------
{template_text}
--------------------

PROFILE INFORMATION:
--------------------
{profile_text}
--------------------

RULES:
1. Follow template structure exactly.
2. Preserve all profile information.
3. Rewrite professionally.
4. Create missing sections only if reasonable.
5. Output only the final formatted profile.
"""

    try:
        response = client.chat.completions.create(
            model=DEPLOYMENT_NAME,
            messages=[
                {
                    "role": "system",
                    "content": "You are a professional executive profile formatter."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.2,
            max_tokens=4000
        )

        return response.choices[0].message.content

    except Exception as e:
        st.error(f"OpenAI request failed: {e}")
        return None


# -----------------------------
# Streamlit UI
# -----------------------------
st.set_page_config(
    page_title="Profile Formatter",
    page_icon="📄",
    layout="wide"
)

st.title("📄 Profile Formatter")

st.write(
    "Upload a raw profile and a template document. "
    "The application will generate a professionally formatted profile "
    "based on the template structure."
)

profile_file = st.file_uploader(
    "Upload Profile Document",
    type=["docx"]
)

template_file = st.file_uploader(
    "Upload Template Document",
    type=["docx"]
)

if st.button("Generate Formatted Profile"):

    if not profile_file:
        st.warning("Please upload a Profile document.")
        st.stop()

    if not template_file:
        st.warning("Please upload a Template document.")
        st.stop()

    with st.spinner("Generating profile..."):

        profile_text = read_docx(profile_file)
        template_text = read_docx(template_file)

        if not profile_text.strip():
            st.error("Profile document is empty.")
            st.stop()

        if not template_text.strip():
            st.error("Template document is empty.")
            st.stop()

        formatted_profile = generate_profile(
            profile_text,
            template_text
        )

    if formatted_profile:

        st.success("Profile generated successfully.")

        st.subheader("Generated Profile")

        st.text_area(
            "Output",
            formatted_profile,
            height=500
        )

        word_file = create_word_document(
            formatted_profile
        )

        st.download_button(
            label="📥 Download Word File",
            data=word_file,
            file_name="Formatted_Profile.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
